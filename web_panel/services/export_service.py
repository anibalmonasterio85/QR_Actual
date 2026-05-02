"""
QR Access Control PRO - Export Service
Generate reports in PDF, Excel, and Word formats.
"""
import os
import io
from datetime import datetime


def export_pdf(logs, title="Reporte de Accesos"):
    """Export access logs to PDF. Returns bytes buffer."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, title, ln=True, align='C')
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 8, f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', ln=True, align='C')
    pdf.ln(5)

    # Table header
    pdf.set_font('Arial', 'B', 9)
    pdf.set_fill_color(102, 126, 234)
    pdf.set_text_color(255, 255, 255)
    col_widths = [15, 40, 50, 25, 45]
    headers = ['ID', 'Usuario', 'QR', 'Resultado', 'Fecha/Hora']
    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], 8, header, 1, 0, 'C', True)
    pdf.ln()

    # Table rows
    pdf.set_font('Arial', '', 8)
    pdf.set_text_color(0, 0, 0)
    for log in logs:
        nombre = log.get('nombre', 'Desconocido') or 'Desconocido'
        qr_text = (log.get('qr_texto', '')[:20] + '...') if log.get('qr_texto', '') else ''
        resultado = log.get('resultado', '')
        fecha = log.get('fecha_hora', '')
        if hasattr(fecha, 'strftime'):
            fecha = fecha.strftime('%Y-%m-%d %H:%M')
        else:
            fecha = str(fecha)[:16]

        if resultado == 'permitido':
            pdf.set_fill_color(220, 255, 220)
        else:
            pdf.set_fill_color(255, 220, 220)

        pdf.cell(col_widths[0], 7, str(log.get('id', '')), 1, 0, 'C', True)
        pdf.cell(col_widths[1], 7, nombre[:18], 1, 0, 'L', True)
        pdf.cell(col_widths[2], 7, qr_text, 1, 0, 'L', True)
        pdf.cell(col_widths[3], 7, resultado.upper(), 1, 0, 'C', True)
        pdf.cell(col_widths[4], 7, fecha, 1, 0, 'C', True)
        pdf.ln()

    # Summary
    pdf.ln(5)
    pdf.set_font('Arial', 'B', 10)
    total = len(logs)
    permitidos = sum(1 for l in logs if l.get('resultado') == 'permitido')
    denegados = total - permitidos
    pdf.cell(0, 8, f'Total: {total} | Permitidos: {permitidos} | Denegados: {denegados}', ln=True)

    return pdf.output(dest='S').encode('latin-1')


def export_excel(logs, title="Reporte de Accesos"):
    """Export access logs to Excel. Returns bytes buffer."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Accesos"

    # Title row
    ws.merge_cells('A1:F1')
    title_cell = ws['A1']
    title_cell.value = title
    title_cell.font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
    title_cell.fill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
    title_cell.alignment = Alignment(horizontal='center')

    # Date row
    ws.merge_cells('A2:F2')
    ws['A2'].value = f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    ws['A2'].font = Font(size=9, italic=True)
    ws['A2'].alignment = Alignment(horizontal='center')

    # Headers
    headers = ['ID', 'Usuario', 'Correo', 'QR Texto', 'Resultado', 'Fecha/Hora']
    header_fill = PatternFill(start_color='764BA2', end_color='764BA2', fill_type='solid')
    header_font = Font(name='Arial', size=10, bold=True, color='FFFFFF')

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')

    # Data rows
    for row_idx, log in enumerate(logs, 5):
        nombre = log.get('nombre', 'Desconocido') or 'Desconocido'
        correo = log.get('correo', '') or ''
        fecha = log.get('fecha_hora', '')
        if hasattr(fecha, 'strftime'):
            fecha = fecha.strftime('%Y-%m-%d %H:%M:%S')

        ws.cell(row=row_idx, column=1, value=log.get('id', ''))
        ws.cell(row=row_idx, column=2, value=nombre)
        ws.cell(row=row_idx, column=3, value=correo)
        ws.cell(row=row_idx, column=4, value=log.get('qr_texto', ''))
        resultado_cell = ws.cell(row=row_idx, column=5, value=log.get('resultado', '').upper())
        ws.cell(row=row_idx, column=6, value=str(fecha))

        if log.get('resultado') == 'permitido':
            resultado_cell.fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
        else:
            resultado_cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')

    # Column widths
    for col_letter, width in [('A', 8), ('B', 25), ('C', 30), ('D', 35), ('E', 15), ('F', 22)]:
        ws.column_dimensions[col_letter].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_word(logs, title="Reporte de Accesos"):
    """Export access logs to Word document. Returns bytes buffer."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph(f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # Summary
    total = len(logs)
    permitidos = sum(1 for l in logs if l.get('resultado') == 'permitido')
    denegados = total - permitidos
    doc.add_paragraph(f'Total de registros: {total} | Permitidos: {permitidos} | Denegados: {denegados}')

    # Table
    if logs:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['ID', 'Usuario', 'Resultado', 'QR', 'Fecha/Hora']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for log in logs:
            row = table.add_row().cells
            row[0].text = str(log.get('id', ''))
            row[1].text = log.get('nombre', 'Desconocido') or 'Desconocido'
            row[2].text = log.get('resultado', '').upper()
            row[3].text = (log.get('qr_texto', '')[:25] + '...') if log.get('qr_texto') else ''
            fecha = log.get('fecha_hora', '')
            row[4].text = fecha.strftime('%Y-%m-%d %H:%M') if hasattr(fecha, 'strftime') else str(fecha)[:16]

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
